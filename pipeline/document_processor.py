"""
SafeAI — Phase 5: Document and Image Processor
-------------------------------------------------
Handles file uploads — images, PDFs, Word docs, and text files.
Extracts all text content before anything reaches the LLM.

Uses PyMuPDF (fitz) for PDFs and images — pure Python,
no system dependencies, works on M1 Mac with macOS 13.

Includes steganography detection on image uploads using
LSB variance analysis — standard digital forensics technique.

What this covers (v1):
    Images      JPG, PNG, GIF, WEBP — text extraction + steg check
    PDFs        Text layer extracted directly via PyMuPDF
    Word docs   Text and tables via python-docx
    Text files  Read directly

What is deferred to v2:
    Scanned PDFs with no text layer
    Audio and video files
    Multilingual documents beyond English
    Handwritten forms
"""

import os
from dataclasses import dataclass, field
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

import fitz

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image as PILImage
    import numpy as np
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ─────────────────────────────────────────────
# Output data structure
# ─────────────────────────────────────────────

@dataclass
class DocumentResult:
    file_path: str
    file_type: str
    extracted_text: str
    image_count: int
    page_count: int
    extraction_method: str
    warnings: list = field(default_factory=list)
    steg_suspected: bool = False
    lsb_variance: float = 1.0

    @property
    def has_content(self):
        return bool(self.extracted_text.strip())

    @property
    def word_count(self):
        return len(self.extracted_text.split())


# ─────────────────────────────────────────────
# Document Processor
# ─────────────────────────────────────────────

class DocumentProcessor:
    """
    Routes file uploads to the appropriate extractor
    based on file type and returns clean text ready
    for the SafeAI pipeline.
    """

    SUPPORTED_IMAGE_TYPES = {
        ".jpg", ".jpeg", ".png", ".gif",
        ".webp", ".bmp", ".tiff"
    }

    # LSB variance below this value suggests steganography.
    # Calibrated for photographs and scanned documents.
    # Digital screenshots score low naturally so we also
    # require a minimum pixel count before flagging.
    STEG_THRESHOLD = 0.05
    STEG_MIN_PIXELS = 10000

    def process(self, file_path: str) -> DocumentResult:
        """
        Main entry point.
        Detects file type and routes to correct processor.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix in self.SUPPORTED_IMAGE_TYPES:
            return self._process_image(path)
        elif suffix == ".pdf":
            return self._process_pdf(path)
        elif suffix in {".docx", ".doc"}:
            return self._process_docx(path)
        elif suffix == ".txt":
            return self._process_txt(path)
        else:
            return DocumentResult(
                file_path=str(path),
                file_type="unsupported",
                extracted_text="",
                image_count=0,
                page_count=0,
                extraction_method="none",
                warnings=[
                    f"Unsupported file type: {suffix}. "
                    f"Supported: images (jpg/png/gif/webp), pdf, docx, txt."
                ]
            )

    def _process_image(self, path: Path) -> DocumentResult:
        """
        Two operations run on every image:

        1. Text extraction via PyMuPDF — works on images
           that have a digital text layer (screenshots,
           typed documents saved as image).
           Does not work on photographs or scanned paper.

        2. Steganography check via LSB variance analysis —
           works on ALL images regardless of content.
           Flags statistical anomalies consistent with
           hidden data encoding.
        """
        warnings = []
        extracted_text = ""
        steg_suspected = False
        lsb_variance = 1.0

        # ── Text extraction ──
        try:
            doc = fitz.open(str(path))
            page = doc[0]
            blocks = page.get_text("blocks")
            text_parts = [
                block[4].strip()
                for block in blocks
                if len(block) >= 5 and block[4].strip()
            ]
            extracted_text = " ".join(text_parts)
            doc.close()

            if not extracted_text.strip():
                warnings.append(
                    "No text layer found in image. "
                    "This image may be a photograph or scanned document. "
                    "Full OCR support for these cases is planned for v2."
                )

        except Exception as e:
            warnings.append(f"Image text extraction error: {str(e)}")

        # ── Steganography check ──
        # Runs regardless of whether text was found.
        # An image with no visible text could still
        # carry hidden data in its pixel values.
        if PIL_AVAILABLE:
            try:
                img = PILImage.open(str(path)).convert("RGB")
                pixels = np.array(img)

                # Extract least significant bits of all channels
                lsb_layer = pixels & 1
                lsb_variance = round(float(np.var(lsb_layer)), 6)

                # Only flag if image is large enough to matter
                # Small images naturally have low variance
                pixel_count = pixels.shape[0] * pixels.shape[1]

                if lsb_variance < self.STEG_THRESHOLD and pixel_count > self.STEG_MIN_PIXELS:
                    steg_suspected = True
                    warnings.append(
                        f"Steganography suspected: LSB variance {lsb_variance} "
                        f"is abnormally low across {pixel_count:,} pixels. "
                        f"Image may contain hidden encoded data. "
                        f"Flagged for compliance review."
                    )

            except Exception:
                # Do not flag on analysis failure
                # Avoid false positives from corrupt images
                lsb_variance = 1.0
                steg_suspected = False
        else:
            warnings.append(
                "Pillow not installed — steganography check skipped. "
                "Run: pip install Pillow"
            )

        return DocumentResult(
            file_path=str(path),
            file_type="image",
            extracted_text=extracted_text,
            image_count=1,
            page_count=1,
            extraction_method="pymupdf_image",
            warnings=warnings,
            steg_suspected=steg_suspected,
            lsb_variance=lsb_variance
        )

    def _process_pdf(self, path: Path) -> DocumentResult:
        """
        Extracts text from PDFs using PyMuPDF.
        Handles multi-page documents.
        Cleans up excessive whitespace from PDF formatting.
        """
        all_text = []
        warnings = []
        page_count = 0
        empty_pages = 0

        try:
            doc = fitz.open(str(path))
            page_count = len(doc)

            for i, page in enumerate(doc):
                page_text = page.get_text("text")
                # Collapse broken whitespace from PDF formatting
                cleaned = " ".join(page_text.split())

                if cleaned.strip():
                    all_text.append(f"[Page {i+1}]\n{cleaned}")
                else:
                    empty_pages += 1

            doc.close()

            if empty_pages > 0:
                warnings.append(
                    f"{empty_pages} page(s) had no extractable text. "
                    f"These may be scanned pages. "
                    f"Full OCR support is planned for v2."
                )

        except Exception as e:
            warnings.append(f"PDF extraction error: {str(e)}")

        return DocumentResult(
            file_path=str(path),
            file_type="pdf",
            extracted_text="\n\n".join(all_text),
            image_count=0,
            page_count=page_count,
            extraction_method="pymupdf_pdf",
            warnings=warnings
        )

    def _process_docx(self, path: Path) -> DocumentResult:
        """
        Extracts text from Word documents.
        Processes all paragraphs and tables.
        """
        if not DOCX_AVAILABLE:
            return DocumentResult(
                file_path=str(path),
                file_type="docx",
                extracted_text="",
                image_count=0,
                page_count=0,
                extraction_method="none",
                warnings=["python-docx not installed. Run: pip install python-docx"]
            )

        all_text = []
        warnings = []

        try:
            doc = Document(str(path))

            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        all_text.append(row_text)

            if not all_text:
                warnings.append(
                    "No text extracted from Word document. "
                    "Document may be empty or image-only."
                )

        except Exception as e:
            warnings.append(f"DOCX extraction error: {str(e)}")

        return DocumentResult(
            file_path=str(path),
            file_type="docx",
            extracted_text="\n".join(all_text),
            image_count=0,
            page_count=1,
            extraction_method="python_docx",
            warnings=warnings
        )

    def _process_txt(self, path: Path) -> DocumentResult:
        """Plain text files — read directly."""
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            return DocumentResult(
                file_path=str(path),
                file_type="txt",
                extracted_text=text,
                image_count=0,
                page_count=1,
                extraction_method="direct_read",
                warnings=[]
            )
        except Exception as e:
            return DocumentResult(
                file_path=str(path),
                file_type="txt",
                extracted_text="",
                image_count=0,
                page_count=0,
                extraction_method="none",
                warnings=[f"Text file read error: {str(e)}"]
            )
